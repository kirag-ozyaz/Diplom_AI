# docx_to_md_images_1.py
# Конвертация DOCX в Markdown с извлечением изображений
# Изображения вставляются в виде ссылок ![alt](path), а не base64
# Нормально выгружает image файлы и вставляются в md файл
# import os
import sys
import base64
from pathlib import Path
import mammoth
from bs4 import BeautifulSoup
import re
import shutil
from zipfile import ZipFile
from docx import Document
import hashlib


def clean_hidden_tags_in_docx(docx_path):
    """
    Удалить скрытые метки вроде:
      #G0...
      #M12291 901873648 ...
      #S
      #X... (любой символ после # + цифры/пробелы)
    из всех параграфов и ячеек таблиц в .docx.
    Возвращает изменённый Document (в памяти).
    """
    doc = Document(docx_path)

    # Шаблон для скрытых меток (включая строки целиком и фрагменты)
    # pattern = re.compile(
    #     r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # полная строка
    #     r'|'
    #     r'(?:\s+)?#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?(?:\s+|#S)?',     # фрагмент внутри строки
    #     re.IGNORECASE
    # )
    pattern = re.compile(
        r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # полная строка (оставляем как есть)
        r'|'
        r'(?:\s+)?(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z])(?:\s+)?',  # ← ИЗМЕНЕНО: добавлен |#[A-Z]
        re.IGNORECASE
    )

    # Очищаем параграфы
    for para in doc.paragraphs:
        if para.text.strip():
            cleaned_text = pattern.sub('', para.text)
            # Убираем лишние пробелы и пустые строки
            cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text).strip()
            if not cleaned_text:
                para.clear()  # полностью очищаем пустой параграф
            else:
                para.text = cleaned_text

    # Очищаем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cleaned_text = pattern.sub('', cell.text)
                    cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text).strip()
                    cell.text = cleaned_text

    return doc


def clean_hidden_tags_in_markdown(markdown_content):
    """
    Удалить скрытые метки вроде:
      #G0...
      #M12291 901873648 ...
      #S
      #X... (любой символ после # + цифры/пробелы)
    из markdown текста.
    """
    pattern = re.compile(
        r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # полная строка
        r'|'
        r'(?:\s+)?(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z])(?:\s+)?',  # фрагмент внутри строки
        re.IGNORECASE | re.MULTILINE
    )
    
    # Очищаем скрытые метки
    cleaned_content = pattern.sub('', markdown_content)
    # Убираем лишние пробелы и пустые строки
    cleaned_content = re.sub(r'\s{3,}', '\n\n', cleaned_content)  # Заменяем 3+ пробела на двойной перенос строки
    cleaned_content = re.sub(r'\n{4,}', '\n\n\n', cleaned_content)  # Ограничиваем множественные переносы строк
    
    return cleaned_content


def merge_split_headers(markdown_content):
    """
    Объединить заголовки, разбитые на несколько строк.
    
    Если заголовок (строка начинающаяся с #) заканчивается без знака препинания,
    а следующая строка не является новым заголовком, списком, таблицей или кодом,
    то объединить их в одну строку. Может объединять несколько строк подряд.
    
    Пример:
        # Первая часть заголовка
        вторая часть заголовка
        третья часть
        
        Становится:
        # Первая часть заголовка вторая часть заголовка третья часть
    """
    lines = markdown_content.split('\n')
    merged_lines = []
    i = 0
    
    def is_block_start(line):
        """Проверить, является ли строка началом нового блока."""
        stripped = line.strip()
        if not stripped:
            return True
        return (re.match(r'^#{1,6}\s+', stripped) or
                re.match(r'^[-*+]\s+', stripped) or
                re.match(r'^\d+\.\s+', stripped) or
                '|' in stripped or
                re.match(r'^```', stripped) or
                re.match(r'^ {4,}', stripped) or
                re.match(r'^[-*]{3,}$', stripped))
    
    while i < len(lines):
        current_line = lines[i]
        
        # Проверяем, является ли текущая строка заголовком
        header_match = re.match(r'^(#{1,6})\s+(.+)$', current_line)
        
        if header_match:
            header_level = header_match.group(1)
            header_parts = [header_match.group(2).rstrip()]
            j = i + 1
            
            # Собираем все последующие строки, которые являются продолжением заголовка
            while j < len(lines):
                next_line = lines[j].strip()
                
                # Если следующая строка пустая или является новым блоком - останавливаемся
                if not next_line or is_block_start(lines[j]):
                    break
                
                # Если текущая последняя часть заканчивается на знак препинания (.!?;:),
                # то следующая строка скорее всего не является продолжением
                if header_parts and header_parts[-1] and header_parts[-1][-1] in '.!?;:':
                    break
                
                # Добавляем следующую строку как продолжение заголовка
                header_parts.append(next_line)
                j += 1
            
            # Объединяем все части заголовка в одну строку
            merged_header = f"{header_level} {' '.join(header_parts)}"
            merged_lines.append(merged_header)
            i = j  # Переходим к строке после последней объединенной
            continue
        
        # Если это не заголовок - просто добавляем
        merged_lines.append(current_line)
        i += 1
    
    return '\n'.join(merged_lines)


def extract_images_and_fix_refs(docx_path, output_dir, file_stem):
    """Извлечь изображения из .docx и вернуть словарь {rId: имя_файла}"""
    # Создаем папку с именем image_<имя_файла>
    images_dir = output_dir / f"image_{file_stem}"
    
    # Очищаем папку, если она существует
    if images_dir.exists():
        shutil.rmtree(images_dir)
    
    images_dir.mkdir(exist_ok=True)

    image_map = {}
    image_counter = 1

    # Извлекаем изображения через ZIP (т.к. .docx — это ZIP-архив)
    with ZipFile(docx_path, 'r') as docx_zip:
        # Сначала читаем relationships
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path not in docx_zip.namelist():
            print(f"  ⚠️  Файл {rels_path} не найден в архиве")
            return image_map
        
        rels_xml = docx_zip.read(rels_path).decode('utf-8')
        soup = BeautifulSoup(rels_xml, 'xml')
        
        # Получаем все Relationship элементы
        relationships = soup.find_all('Relationship')
        print(f"  📋 Найдено relationships: {len(relationships)}")
        
        # Парсим все связи изображений
        for rel in relationships:
            rel_type = rel.get('Type', '')
            # Проверяем, является ли это изображением
            if 'image' not in rel_type.lower():
                continue
            
            r_id = rel.get('Id', '')
            target = rel.get('Target', '')
            
            if not r_id or not target:
                print(f"  ⚠️  Пропущена связь: Id={r_id}, Target={target}")
                continue
            
            # Формируем полный путь к изображению
            # Target может быть: "media/image1.png" или "../media/image1.png" или просто "image1.png"
            if target.startswith('media/'):
                img_path_in_zip = f'word/{target}'
            elif target.startswith('../media/'):
                img_path_in_zip = f'word/{target[3:]}'  # Убираем ../
            elif '/' in target:
                # Если есть слэш, но не media/, пробуем как есть
                img_path_in_zip = f'word/{target}' if not target.startswith('word/') else target
            else:
                # Просто имя файла
                img_path_in_zip = f'word/media/{target}'
            
            # Проверяем, существует ли файл
            if img_path_in_zip not in docx_zip.namelist():
                # Пробуем альтернативные варианты
                alt_paths = [
                    f'word/media/{Path(target).name}',
                    f'word/{target}',
                    target
                ]
                found = False
                for alt_path in alt_paths:
                    if alt_path in docx_zip.namelist():
                        img_path_in_zip = alt_path
                        found = True
                        break
                
                if not found:
                    print(f"  ❌ Изображение не найдено: rId={r_id}, target={target}, пробовали: {img_path_in_zip}")
                    # Выводим список всех файлов в архиве для отладки
                    media_files = [f for f in docx_zip.namelist() if 'media' in f.lower() or 'image' in f.lower()]
                    if media_files:
                        print(f"     Доступные медиа-файлы: {media_files[:5]}...")  # Показываем первые 5
                    continue
            
            # Используем оригинальное имя файла из target
            # Извлекаем имя файла из пути (например, "media/image1.png" -> "image1.png")
            original_name = Path(target).name
            
            # Если имя файла пустое или нет расширения, используем счетчик
            if not original_name or not Path(original_name).suffix:
                ext = Path(target).suffix.lower()
                if not ext or ext not in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'):
                    ext = '.png'
                original_name = f"image{image_counter:03d}{ext}"
            
            # Проверяем, не существует ли уже файл с таким именем
            img_name = original_name
            img_path = images_dir / img_name
            counter = 1
            while img_path.exists():
                # Если файл уже существует, добавляем суффикс
                stem = Path(original_name).stem
                ext = Path(original_name).suffix
                img_name = f"{stem}_{counter}{ext}"
                img_path = images_dir / img_name
                counter += 1
            
            try:
                with open(img_path, 'wb') as f:
                    f.write(docx_zip.read(img_path_in_zip))
                
                image_map[r_id] = img_name
                print(f"  ✓ {r_id} → {img_name} (из {img_path_in_zip})")
                image_counter += 1
            except Exception as e:
                print(f"  ❌ Ошибка при сохранении изображения {r_id}: {e}")

    return image_map


def replace_image_tags_in_html(html, image_map, images_folder_name, images_dir):
    """Заменить <img> теги на корректные пути, сохраняя все изображения"""
    soup = BeautifulSoup(html, 'html.parser')
    #image_counter = image_counter_start

    for img in soup.find_all('img'):
        src = img.get('src', '')
        new_src = None

        # print('replace ', src[:80])
        # 1. Обработка rId (ПЕРВЫМ, чтобы обработать все изображения из image_map)
        if src.startswith('rId') or (match := re.search(r'(rId\d+)', src)):
            r_id = match.group(1) if match else src
            if r_id in image_map:
                new_src = f"{images_folder_name}/{image_map[r_id]}"
                print(f"  ✓ Обработан rId: {r_id} → {image_map[r_id]}")

        # 2. Обработка base64 (только если rId не найден)
        elif src.startswith('data:image/'):
            try:
                match = re.match(r'data:image/(\w+);base64,(.+)', src)
                if match:
                    img_format = match.group(1).lower()
                    base64_data = match.group(2)
                    ext = {'png': '.png', 'jpeg': '.jpg', 'jpg': '.jpg', 'gif': '.gif', 'bmp': '.bmp',
                           'webp': '.webp'}.get(img_format, '.png')

                    img_data = base64.b64decode(base64_data)
                    import hashlib
                    img_hash = hashlib.md5(img_data).hexdigest()[:8]
                    base_stem = f"img_{img_hash}"
                    
                    # Проверяем существование файла и добавляем счетчик только если нужно
                    img_name = f"{base_stem}{ext}"
                    img_path = images_dir / img_name
                    counter = 1
                    while img_path.exists():
                        # Всегда используем базовое имя с новым счетчиком
                        img_name = f"{base_stem}_{counter}{ext}"
                        img_path = images_dir / img_name
                        counter += 1
                        # Защита от бесконечного цикла и слишком длинных имен
                        if counter > 10000:
                            # Если слишком много попыток, используем timestamp
                            import time
                            img_name = f"{base_stem}_{int(time.time())}{ext}"
                            img_path = images_dir / img_name
                            break

                    with open(img_path, 'wb') as f:
                        f.write(img_data)

                    new_src = f"{images_folder_name}/{img_name}"
                    print(f"  ✓ Извлечено из base64: {img_name}")
            except Exception as e:
                print(f"  ⚠️ Ошибка base64: {e}")

        # 3. Если не распознано — НЕ удаляем, оставляем для отладки
        if new_src:
            img['src'] = new_src
            if not img.get('alt'):
                img['alt'] = Path(new_src).name
        else:
            # Оставляем тег, но помечаем для отладки
            print(f"  ⚠️ Необработанное изображение (src='{src[:60]}...'), оставлено как есть")

    return str(soup)


def fix_images_in_markdown(markdown_content, images_dir, images_folder_name):
    """
    Заменяет base64 изображения и оставшиеся <img> теги на ссылки вида ![alt](path)
    """
    # 1. Обработка base64 изображений: ![](data:image/...)
    def replace_base64_images(text):
        def process_base64(match):
            base64_data = match.group(1)
            try:
                data_match = re.match(r'data:image/(\w+);base64,(.+)', base64_data)
                if not data_match:
                    return match.group(0)
                
                img_format = data_match.group(1).lower()
                base64_str = data_match.group(2)
                
                ext_map = {'png': '.png', 'jpeg': '.jpg', 'jpg': '.jpg', 
                          'gif': '.gif', 'bmp': '.bmp', 'webp': '.webp'}
                ext = ext_map.get(img_format, '.png')
                
                img_data = base64.b64decode(base64_str)
                img_hash = hashlib.md5(img_data).hexdigest()[:8]
                img_name = f"img_{img_hash}{ext}"
                img_path = images_dir / img_name
                
                if not img_path.exists():
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                
                return f'![image]({images_folder_name}/{img_name})'
            except Exception as e:
                print(f"  ⚠️ Ошибка при обработке base64: {e}")
                return match.group(0)
        
        pattern = r'!\[\]\((data:image/[^)]+)\)'
        return re.sub(pattern, process_base64, text)
    
    # 2. Обработка оставшихся <img> тегов
    def replace_img_tags(text):
        def process_img_tag(match):
            img_tag = match.group(0)
            soup = BeautifulSoup(img_tag, 'html.parser')
            img = soup.find('img')
            if not img:
                return img_tag
            
            src = img.get('src', '')
            alt = img.get('alt', Path(src).name if src else 'image')
            
            if src and not src.startswith(('data:', 'http://', 'https://')):
                return f'![{alt}]({src})'
            
            return img_tag
        
        pattern = r'<img[^>]+>'
        return re.sub(pattern, process_img_tag, text, flags=re.IGNORECASE)
    
    markdown_content = replace_base64_images(markdown_content)
    markdown_content = replace_img_tags(markdown_content)
    
    return markdown_content

def docx_to_md_with_images(docx_path, output_dir=None, merge_headers=False):
    """
    Конвертирует DOCX в Markdown с извлечением изображений.
    
    Args:
        docx_path: Путь к файлу .docx
        output_dir: Директория для сохранения изображений (по умолчанию - папка с исходным файлом)
        merge_headers: Если True, автоматически объединяет заголовки, разбитые на несколько строк
    """
    docx_path = Path(docx_path).resolve()
    if docx_path.suffix.lower() != '.docx':
        print("❌ Поддерживается только .docx")
        sys.exit(1)

    # Если output_dir не указан, используем папку с исходным файлом
    if output_dir is None:
        output_dir = docx_path.parent
    else:
        output_dir = Path(output_dir)
    
    # md_path = output_dir / f"{docx_path.stem}.md"

    # Шаг 1: Извлекаем изображения (из оригинального файла, как в docx_to_md_images_3.py)
    print("🖼️  Извлечение изображений...")
    file_stem = docx_path.stem
    images_folder_name = f"image_{file_stem}"
    image_map = extract_images_and_fix_refs(docx_path, output_dir, file_stem)
    print(f"  Найдено изображений: {len(image_map)}")

    # Шаг 2: Конвертируем в HTML через mammoth
    # Используем ОРИГИНАЛЬНЫЙ файл, чтобы mammoth оставил изображения как rId (не преобразовывал в base64)
    print("🔄 Конвертация в HTML...")
    
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value

    # Шаг 3: Заменяем ссылки на изображения
    print("🔗 Проверка ссылок на изображения...")
    # Отладочный вывод: проверяем, какие img теги есть в HTML
    soup_debug = BeautifulSoup(html, 'html.parser')
    img_tags = soup_debug.find_all('img')
    if img_tags:
        print(f"  📋 Найдено img тегов в HTML (первые три тэга): {len(img_tags)}")
        for i, img in enumerate(img_tags[:3]):  # Показываем первые 3 для отладки
            src = img.get('src', '')
            print(f"     img[{i}]: src='{src[:80]}...' (первые 80 символов)")
    
    # Получаем путь к папке с изображениями
    images_dir = output_dir / images_folder_name

    
    html = replace_image_tags_in_html(html, image_map, images_folder_name, images_dir)
    
    # Проверяем результат замены ссылок
    soup_after = BeautifulSoup(html, 'html.parser')
    img_tags_after = soup_after.find_all('img')
    if img_tags_after:
        print(f"  📋 После замены найдено img тегов: {len(img_tags_after)}")
        for i, img in enumerate(img_tags_after[:3]):  # Показываем первые 3 для отладки
            src = img.get('src', '')
            alt = img.get('alt', '')
            print(f"     img[{i}]: src='{src}', alt='{alt}'")

    # Шаг 4: Конвертируем HTML → Markdown
    print("📝 Преобразование в Markdown...")
    from markdownify import markdownify as md

    markdown_content = md(html, heading_style="ATX", strip=['style'])

    # Шаг 4.5: Исправляем изображения (заменяем base64 и <img> на ссылки)
    print("🔧 Исправление ссылок на изображения...")
    markdown_content = fix_images_in_markdown(markdown_content, images_dir, images_folder_name)

    # Шаг 4.6: Очищаем скрытые метки из markdown
    print("🧹 Очистка скрытых меток...")
    markdown_content = clean_hidden_tags_in_markdown(markdown_content)

    # Шаг 4.7: Объединяем разбитые заголовки (если включено)
    if merge_headers:
        print("🔗 Объединение разбитых заголовков...")
        markdown_content = merge_split_headers(markdown_content)

    # # Шаг 5: Сохраняем
    # with open(md_path, 'w', encoding='utf-8') as f:
    #     f.write(markdown_content)

    print(f"\n✅ Готово!")
    # print(f"📄 Markdown: {md_path}")
    print(f"🖼️  Изображения: {output_dir / images_folder_name}")
    
    return markdown_content

# file_docx = r"X:\Учеба_УИИ\Итоговы_Проект\Этап №2.  AI_ML  Сбор базы\Нормативная база\ПУЭ\DOCX\2.5.docx"
# file_docx = r"X:\Учеба_УИИ\Итоговы_Проект\data\raw\Нормативная база\ПУЭ\DOCX\2.5.docx"
# output_file_dir = r"X:\Учеба_УИИ\Итоговы_Проект\data\extracted"

ROOT = Path(__file__).resolve().parent.parent.parent.parent
output_file_dir = ROOT / "data" / "extracted"
file_docx = ROOT / "data" / "raw"/"Нормативная база"/"ПУЭ"/"DOCX" / "7.5.docx"

if __name__ == "__main__":
    sys_argv = None
    output_dir_arg = None
    try:
        # Рабочая версия: проверяем наличие аргументов командной строки
        if len(sys.argv) == 2 and sys.argv[1].strip():
            sys_argv = sys.argv[1]
        elif len(sys.argv) == 3 and sys.argv[1].strip() and sys.argv[2].strip():
            sys_argv = sys.argv[1]
            output_dir_arg = sys.argv[2]
        else:
            # Для тестов: используем file_docx
            sys_argv = file_docx
            output_dir_arg = output_file_dir
        
        # Проверяем существование файла
        docx_path = Path(sys_argv)
        if not docx_path.exists():
            raise FileNotFoundError(f"Файл не найден: {sys_argv}")
        
        # Определяем output_dir
        if output_dir_arg:
            output_dir = Path(output_dir_arg)
            # Проверяем существование директории
            if not output_dir.exists():
                raise FileNotFoundError(f"Директория не найдена: {output_dir_arg}")
            if not output_dir.is_dir():
                raise NotADirectoryError(f"Указанный путь не является директорией: {output_dir_arg}")
        else:
            output_dir = docx_path.parent
        
        md_path = output_dir / f"{docx_path.stem}.md"

        markdown_content = docx_to_md_with_images(sys_argv, output_dir)

        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        print(f"📄 Markdown: {md_path}")
    except FileNotFoundError as e:
        print(f"❌ Ошибка: {e}")
        print("Использование: python docx_to_md_with_images.py <файл.docx> [output_dir]")
        sys.exit(1)
    except NotADirectoryError as e:
        print(f"❌ Ошибка: {e}")
        print("Использование: python docx_to_md_with_images.py <файл.docx> [output_dir]")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        print("Использование: python docx_to_md_with_images.py <файл.docx> [output_dir]")
        sys.exit(1)