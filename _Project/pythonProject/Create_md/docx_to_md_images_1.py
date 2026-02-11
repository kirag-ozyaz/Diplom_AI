# import os
import sys
import base64
from pathlib import Path
import mammoth
from bs4 import BeautifulSoup
import re
import shutil
from zipfile import ZipFile
from docx import Document
import hashlib


def clean_hidden_tags_in_docx(docx_path):
    """
    –£–¥–∞–ª—è–µ—Ç —Å–∫—Ä—ã—Ç—ã–µ –º–µ—Ç–∫–∏ –≤—Ä–æ–¥–µ:
      #G0...
      #M12291 901873648 ...
      #S
      #X... (–ª—é–±–æ–π —Å–∏–º–≤–æ–ª –ø–æ—Å–ª–µ # + —Ü–∏—Ñ—Ä—ã/–ø—Ä–æ–±–µ–ª—ã)
    –∏–∑ –≤—Å–µ—Ö –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –∏ —è—á–µ–µ–∫ —Ç–∞–±–ª–∏—Ü –≤ .docx.
    –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç –∏–∑–º–µ–Ω—ë–Ω–Ω—ã–π Document (–≤ –ø–∞–º—è—Ç–∏).
    """
    doc = Document(docx_path)

    # –®–∞–±–ª–æ–Ω –¥–ª—è —Å–∫—Ä—ã—Ç—ã—Ö –º–µ—Ç–æ–∫ (–≤–∫–ª—é—á–∞—è —Å—Ç—Ä–æ–∫–∏ —Ü–µ–ª–∏–∫–æ–º –∏ —Ñ—Ä–∞–≥–º–µ–Ω—Ç—ã)
    # pattern = re.compile(
    #     r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # –ø–æ–ª–Ω–∞—è —Å—Ç—Ä–æ–∫–∞
    #     r'|'
    #     r'(?:\s+)?#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?(?:\s+|#S)?',     # —Ñ—Ä–∞–≥–º–µ–Ω—Ç –≤–Ω—É—Ç—Ä–∏ —Å—Ç—Ä–æ–∫–∏
    #     re.IGNORECASE
    # )
    pattern = re.compile(
        r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # –ø–æ–ª–Ω–∞—è —Å—Ç—Ä–æ–∫–∞ (–æ—Å—Ç–∞–≤–ª—è–µ–º –∫–∞–∫ –µ—Å—Ç—å)
        r'|'
        r'(?:\s+)?(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z])(?:\s+)?',  # ‚Üê –ò–ó–ú–ï–ù–ï–ù–û: –¥–æ–±–∞–≤–ª–µ–Ω |#[A-Z]
        re.IGNORECASE
    )

    # –û—á–∏—â–∞–µ–º –ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã
    for para in doc.paragraphs:
        if para.text.strip():
            cleaned_text = pattern.sub('', para.text)
            # –£–±–∏—Ä–∞–µ–º –ª–∏—à–Ω–∏–µ –ø—Ä–æ–±–µ–ª—ã –∏ –ø—É—Å—Ç—ã–µ —Å—Ç—Ä–æ–∫–∏
            cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text).strip()
            if not cleaned_text:
                para.clear()  # –ø–æ–ª–Ω–æ—Å—Ç—å—é –æ—á–∏—â–∞–µ–º –ø—É—Å—Ç–æ–π –ø–∞—Ä–∞–≥—Ä–∞—Ñ
            else:
                para.text = cleaned_text

    # –û—á–∏—â–∞–µ–º —Ç–∞–±–ª–∏—Ü—ã
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cleaned_text = pattern.sub('', cell.text)
                    cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text).strip()
                    cell.text = cleaned_text

    return doc


def clean_hidden_tags_in_markdown(markdown_content):
    """
    –£–¥–∞–ª—è–µ—Ç —Å–∫—Ä—ã—Ç—ã–µ –º–µ—Ç–∫–∏ –≤—Ä–æ–¥–µ:
      #G0...
      #M12291 901873648 ...
      #S
      #X... (–ª—é–±–æ–π —Å–∏–º–≤–æ–ª –ø–æ—Å–ª–µ # + —Ü–∏—Ñ—Ä—ã/–ø—Ä–æ–±–µ–ª—ã)
    –∏–∑ markdown —Ç–µ–∫—Å—Ç–∞.
    """
    pattern = re.compile(
        r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # –ø–æ–ª–Ω–∞—è —Å—Ç—Ä–æ–∫–∞
        r'|'
        r'(?:\s+)?(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z])(?:\s+)?',  # —Ñ—Ä–∞–≥–º–µ–Ω—Ç –≤–Ω—É—Ç—Ä–∏ —Å—Ç—Ä–æ–∫–∏
        re.IGNORECASE | re.MULTILINE
    )
    
    # –û—á–∏—â–∞–µ–º —Å–∫—Ä—ã—Ç—ã–µ –º–µ—Ç–∫–∏
    cleaned_content = pattern.sub('', markdown_content)
    # –£–±–∏—Ä–∞–µ–º –ª–∏—à–Ω–∏–µ –ø—Ä–æ–±–µ–ª—ã –∏ –ø—É—Å—Ç—ã–µ —Å—Ç—Ä–æ–∫–∏
    cleaned_content = re.sub(r'\s{3,}', '\n\n', cleaned_content)  # –ó–∞–º–µ–Ω—è–µ–º 3+ –ø—Ä–æ–±–µ–ª–∞ –Ω–∞ –¥–≤–æ–π–Ω–æ–π –ø–µ—Ä–µ–Ω–æ—Å —Å—Ç—Ä–æ–∫–∏
    cleaned_content = re.sub(r'\n{4,}', '\n\n\n', cleaned_content)  # –û–≥—Ä–∞–Ω–∏—á–∏–≤–∞–µ–º –º–Ω–æ–∂–µ—Å—Ç–≤–µ–Ω–Ω—ã–µ –ø–µ—Ä–µ–Ω–æ—Å—ã —Å—Ç—Ä–æ–∫
    
    return cleaned_content


def extract_images_and_fix_refs(docx_path, output_dir, file_stem):
    """–ò–∑–≤–ª–µ–∫–∞–µ—Ç –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∏–∑ .docx –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Å–ª–æ–≤–∞—Ä—å {rId: –∏–º—è_—Ñ–∞–π–ª–∞}"""
    # –°–æ–∑–¥–∞–µ–º –ø–∞–ø–∫—É —Å –∏–º–µ–Ω–µ–º image_<–∏–º—è_—Ñ–∞–π–ª–∞>
    images_dir = output_dir / f"image_{file_stem}"
    
    # –û—á–∏—â–∞–µ–º –ø–∞–ø–∫—É, –µ—Å–ª–∏ –æ–Ω–∞ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç
    if images_dir.exists():
        shutil.rmtree(images_dir)
    
    images_dir.mkdir(exist_ok=True)

    image_map = {}
    image_counter = 1

    # –ò–∑–≤–ª–µ–∫–∞–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è —á–µ—Ä–µ–∑ ZIP (—Ç.–∫. .docx ‚Äî —ç—Ç–æ ZIP-–∞—Ä—Ö–∏–≤)
    with ZipFile(docx_path, 'r') as docx_zip:
        # –°–Ω–∞—á–∞–ª–∞ —á–∏—Ç–∞–µ–º relationships
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path not in docx_zip.namelist():
            print(f"  ‚ö†Ô∏è  –§–∞–π–ª {rels_path} –Ω–µ –Ω–∞–π–¥–µ–Ω –≤ –∞—Ä—Ö–∏–≤–µ")
            return image_map
        
        rels_xml = docx_zip.read(rels_path).decode('utf-8')
        soup = BeautifulSoup(rels_xml, 'xml')
        
        # –ü–æ–ª—É—á–∞–µ–º –≤—Å–µ Relationship —ç–ª–µ–º–µ–Ω—Ç—ã
        relationships = soup.find_all('Relationship')
        print(f"  üìã –ù–∞–π–¥–µ–Ω–æ relationships: {len(relationships)}")
        
        # –ü–∞—Ä—Å–∏–º –≤—Å–µ —Å–≤—è–∑–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π
        for rel in relationships:
            rel_type = rel.get('Type', '')
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ —ç—Ç–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ–º
            if 'image' not in rel_type.lower():
                continue
            
            r_id = rel.get('Id', '')
            target = rel.get('Target', '')
            
            if not r_id or not target:
                print(f"  ‚ö†Ô∏è  –ü—Ä–æ–ø—É—â–µ–Ω–∞ —Å–≤—è–∑—å: Id={r_id}, Target={target}")
                continue
            
            # –§–æ—Ä–º–∏—Ä—É–µ–º –ø–æ–ª–Ω—ã–π –ø—É—Ç—å –∫ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—é
            # Target –º–æ–∂–µ—Ç –±—ã—Ç—å: "media/image1.png" –∏–ª–∏ "../media/image1.png" –∏–ª–∏ –ø—Ä–æ—Å—Ç–æ "image1.png"
            if target.startswith('media/'):
                img_path_in_zip = f'word/{target}'
            elif target.startswith('../media/'):
                img_path_in_zip = f'word/{target[3:]}'  # –£–±–∏—Ä–∞–µ–º ../
            elif '/' in target:
                # –ï—Å–ª–∏ –µ—Å—Ç—å —Å–ª—ç—à, –Ω–æ –Ω–µ media/, –ø—Ä–æ–±—É–µ–º –∫–∞–∫ –µ—Å—Ç—å
                img_path_in_zip = f'word/{target}' if not target.startswith('word/') else target
            else:
                # –ü—Ä–æ—Å—Ç–æ –∏–º—è —Ñ–∞–π–ª–∞
                img_path_in_zip = f'word/media/{target}'
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —Å—É—â–µ—Å—Ç–≤—É–µ—Ç –ª–∏ —Ñ–∞–π–ª
            if img_path_in_zip not in docx_zip.namelist():
                # –ü—Ä–æ–±—É–µ–º –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—ã–µ –≤–∞—Ä–∏–∞–Ω—Ç—ã
                alt_paths = [
                    f'word/media/{Path(target).name}',
                    f'word/{target}',
                    target
                ]
                found = False
                for alt_path in alt_paths:
                    if alt_path in docx_zip.namelist():
                        img_path_in_zip = alt_path
                        found = True
                        break
                
                if not found:
                    print(f"  ‚ùå –ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ: rId={r_id}, target={target}, –ø—Ä–æ–±–æ–≤–∞–ª–∏: {img_path_in_zip}")
                    # –í—ã–≤–æ–¥–∏–º —Å–ø–∏—Å–æ–∫ –≤—Å–µ—Ö —Ñ–∞–π–ª–æ–≤ –≤ –∞—Ä—Ö–∏–≤–µ –¥–ª—è –æ—Ç–ª–∞–¥–∫–∏
                    media_files = [f for f in docx_zip.namelist() if 'media' in f.lower() or 'image' in f.lower()]
                    if media_files:
                        print(f"     –î–æ—Å—Ç—É–ø–Ω—ã–µ –º–µ–¥–∏–∞-—Ñ–∞–π–ª—ã: {media_files[:5]}...")  # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø–µ—Ä–≤—ã–µ 5
                    continue
            
            # –ò—Å–ø–æ–ª—å–∑—É–µ–º –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–µ –∏–º—è —Ñ–∞–π–ª–∞ –∏–∑ target
            # –ò–∑–≤–ª–µ–∫–∞–µ–º –∏–º—è —Ñ–∞–π–ª–∞ –∏–∑ –ø—É—Ç–∏ (–Ω–∞–ø—Ä–∏–º–µ—Ä, "media/image1.png" -> "image1.png")
            original_name = Path(target).name
            
            # –ï—Å–ª–∏ –∏–º—è —Ñ–∞–π–ª–∞ –ø—É—Å—Ç–æ–µ –∏–ª–∏ –Ω–µ—Ç —Ä–∞—Å—à–∏—Ä–µ–Ω–∏—è, –∏—Å–ø–æ–ª—å–∑—É–µ–º —Å—á–µ—Ç—á–∏–∫
            if not original_name or not Path(original_name).suffix:
                ext = Path(target).suffix.lower()
                if not ext or ext not in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'):
                    ext = '.png'
                original_name = f"image{image_counter:03d}{ext}"
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç –ª–∏ —É–∂–µ —Ñ–∞–π–ª —Å —Ç–∞–∫–∏–º –∏–º–µ–Ω–µ–º
            img_name = original_name
            img_path = images_dir / img_name
            counter = 1
            while img_path.exists():
                # –ï—Å–ª–∏ —Ñ–∞–π–ª —É–∂–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç, –¥–æ–±–∞–≤–ª—è–µ–º —Å—É—Ñ—Ñ–∏–∫—Å
                stem = Path(original_name).stem
                ext = Path(original_name).suffix
                img_name = f"{stem}_{counter}{ext}"
                img_path = images_dir / img_name
                counter += 1
            
            try:
                with open(img_path, 'wb') as f:
                    f.write(docx_zip.read(img_path_in_zip))
                
                image_map[r_id] = img_name
                print(f"  ‚úì {r_id} ‚Üí {img_name} (–∏–∑ {img_path_in_zip})")
                image_counter += 1
            except Exception as e:
                print(f"  ‚ùå –û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è {r_id}: {e}")

    return image_map


def replace_image_tags_in_html(html, image_map, images_folder_name, images_dir):
    """–ó–∞–º–µ–Ω—è–µ—Ç <img> —Ç–µ–≥–∏ –Ω–∞ –∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –ø—É—Ç–∏, —Å–æ—Ö—Ä–∞–Ω—è—è –≤—Å–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è"""
    soup = BeautifulSoup(html, 'html.parser')
    #image_counter = image_counter_start

    for img in soup.find_all('img'):
        src = img.get('src', '')
        new_src = None

        # print('replace ', src[:80])
        # 1. –û–±—Ä–∞–±–æ—Ç–∫–∞ rId (–ü–ï–†–í–´–ú, —á—Ç–æ–±—ã –æ–±—Ä–∞–±–æ—Ç–∞—Ç—å –≤—Å–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∏–∑ image_map)
        if src.startswith('rId') or (match := re.search(r'(rId\d+)', src)):
            r_id = match.group(1) if match else src
            if r_id in image_map:
                new_src = f"{images_folder_name}/{image_map[r_id]}"
                print(f"  ‚úì –û–±—Ä–∞–±–æ—Ç–∞–Ω rId: {r_id} ‚Üí {image_map[r_id]}")

        # 2. –û–±—Ä–∞–±–æ—Ç–∫–∞ base64 (—Ç–æ–ª—å–∫–æ –µ—Å–ª–∏ rId –Ω–µ –Ω–∞–π–¥–µ–Ω)
        elif src.startswith('data:image/'):
            try:
                match = re.match(r'data:image/(\w+);base64,(.+)', src)
                if match:
                    img_format = match.group(1).lower()
                    base64_data = match.group(2)
                    ext = {'png': '.png', 'jpeg': '.jpg', 'jpg': '.jpg', 'gif': '.gif', 'bmp': '.bmp',
                           'webp': '.webp'}.get(img_format, '.png')

                    img_data = base64.b64decode(base64_data)
                    import hashlib
                    img_hash = hashlib.md5(img_data).hexdigest()[:8]
                    base_stem = f"img_{img_hash}"
                    
                    # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å—É—â–µ—Å—Ç–≤–æ–≤–∞–Ω–∏–µ —Ñ–∞–π–ª–∞ –∏ –¥–æ–±–∞–≤–ª—è–µ–º —Å—á–µ—Ç—á–∏–∫ —Ç–æ–ª—å–∫–æ –µ—Å–ª–∏ –Ω—É–∂–Ω–æ
                    img_name = f"{base_stem}{ext}"
                    img_path = images_dir / img_name
                    counter = 1
                    while img_path.exists():
                        # –í—Å–µ–≥–¥–∞ –∏—Å–ø–æ–ª—å–∑—É–µ–º –±–∞–∑–æ–≤–æ–µ –∏–º—è —Å –Ω–æ–≤—ã–º —Å—á–µ—Ç—á–∏–∫–æ–º
                        img_name = f"{base_stem}_{counter}{ext}"
                        img_path = images_dir / img_name
                        counter += 1
                        # –ó–∞—â–∏—Ç–∞ –æ—Ç –±–µ—Å–∫–æ–Ω–µ—á–Ω–æ–≥–æ —Ü–∏–∫–ª–∞ –∏ —Å–ª–∏—à–∫–æ–º –¥–ª–∏–Ω–Ω—ã—Ö –∏–º–µ–Ω
                        if counter > 10000:
                            # –ï—Å–ª–∏ —Å–ª–∏—à–∫–æ–º –º–Ω–æ–≥–æ –ø–æ–ø—ã—Ç–æ–∫, –∏—Å–ø–æ–ª—å–∑—É–µ–º timestamp
                            import time
                            img_name = f"{base_stem}_{int(time.time())}{ext}"
                            img_path = images_dir / img_name
                            break

                    with open(img_path, 'wb') as f:
                        f.write(img_data)

                    new_src = f"{images_folder_name}/{img_name}"
                    print(f"  ‚úì –ò–∑–≤–ª–µ—á–µ–Ω–æ –∏–∑ base64: {img_name}")
            except Exception as e:
                print(f"  ‚ö†Ô∏è –û—à–∏–±–∫–∞ base64: {e}")

        # 3. –ï—Å–ª–∏ –Ω–µ —Ä–∞—Å–ø–æ–∑–Ω–∞–Ω–æ ‚Äî –ù–ï —É–¥–∞–ª—è–µ–º, –æ—Å—Ç–∞–≤–ª—è–µ–º –¥–ª—è –æ—Ç–ª–∞–¥–∫–∏
        if new_src:
            img['src'] = new_src
            if not img.get('alt'):
                img['alt'] = Path(new_src).name
        else:
            # –û—Å—Ç–∞–≤–ª—è–µ–º —Ç–µ–≥, –Ω–æ –ø–æ–º–µ—á–∞–µ–º –¥–ª—è –æ—Ç–ª–∞–¥–∫–∏
            print(f"  ‚ö†Ô∏è –ù–µ–æ–±—Ä–∞–±–æ—Ç–∞–Ω–Ω–æ–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ (src='{src[:60]}...'), –æ—Å—Ç–∞–≤–ª–µ–Ω–æ –∫–∞–∫ –µ—Å—Ç—å")

    return str(soup)


def fix_images_in_markdown(markdown_content, images_dir, images_folder_name):
    """
    –ó–∞–º–µ–Ω—è–µ—Ç base64 –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∏ –æ—Å—Ç–∞–≤—à–∏–µ—Å—è <img> —Ç–µ–≥–∏ –Ω–∞ —Å—Å—ã–ª–∫–∏ –≤–∏–¥–∞ ![alt](path)
    """
    # 1. –û–±—Ä–∞–±–æ—Ç–∫–∞ base64 –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: ![](data:image/...)
    def replace_base64_images(text):
        def process_base64(match):
            base64_data = match.group(1)
            try:
                data_match = re.match(r'data:image/(\w+);base64,(.+)', base64_data)
                if not data_match:
                    return match.group(0)
                
                img_format = data_match.group(1).lower()
                base64_str = data_match.group(2)
                
                ext_map = {'png': '.png', 'jpeg': '.jpg', 'jpg': '.jpg', 
                          'gif': '.gif', 'bmp': '.bmp', 'webp': '.webp'}
                ext = ext_map.get(img_format, '.png')
                
                img_data = base64.b64decode(base64_str)
                img_hash = hashlib.md5(img_data).hexdigest()[:8]
                img_name = f"img_{img_hash}{ext}"
                img_path = images_dir / img_name
                
                if not img_path.exists():
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                
                return f'![image]({images_folder_name}/{img_name})'
            except Exception as e:
                print(f"  ‚ö†Ô∏è –û—à–∏–±–∫–∞ –ø—Ä–∏ –æ–±—Ä–∞–±–æ—Ç–∫–µ base64: {e}")
                return match.group(0)
        
        pattern = r'!\[\]\((data:image/[^)]+)\)'
        return re.sub(pattern, process_base64, text)
    
    # 2. –û–±—Ä–∞–±–æ—Ç–∫–∞ –æ—Å—Ç–∞–≤—à–∏—Ö—Å—è <img> —Ç–µ–≥–æ–≤
    def replace_img_tags(text):
        def process_img_tag(match):
            img_tag = match.group(0)
            soup = BeautifulSoup(img_tag, 'html.parser')
            img = soup.find('img')
            if not img:
                return img_tag
            
            src = img.get('src', '')
            alt = img.get('alt', Path(src).name if src else 'image')
            
            if src and not src.startswith(('data:', 'http://', 'https://')):
                return f'![{alt}]({src})'
            
            return img_tag
        
        pattern = r'<img[^>]+>'
        return re.sub(pattern, process_img_tag, text, flags=re.IGNORECASE)
    
    markdown_content = replace_base64_images(markdown_content)
    markdown_content = replace_img_tags(markdown_content)
    
    return markdown_content

def docx_to_md_with_images(docx_path, output_dir=None):
    docx_path = Path(docx_path).resolve()
    if docx_path.suffix.lower() != '.docx':
        print("‚ùå –ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ—Ç—Å—è —Ç–æ–ª—å–∫–æ .docx")
        sys.exit(1)

    # –ï—Å–ª–∏ output_dir –Ω–µ —É–∫–∞–∑–∞–Ω, –∏—Å–ø–æ–ª—å–∑—É–µ–º –ø–∞–ø–∫—É —Å –∏—Å—Ö–æ–¥–Ω—ã–º —Ñ–∞–π–ª–æ–º
    if output_dir is None:
        output_dir = docx_path.parent
    else:
        output_dir = Path(output_dir)
    
    # md_path = output_dir / f"{docx_path.stem}.md"

    # –®–∞–≥ 1: –ò–∑–≤–ª–µ–∫–∞–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è (–∏–∑ –æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–æ–≥–æ —Ñ–∞–π–ª–∞, –∫–∞–∫ –≤ docx_to_md_images_3.py)
    print("üñºÔ∏è  –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π...")
    file_stem = docx_path.stem
    images_folder_name = f"image_{file_stem}"
    image_map = extract_images_and_fix_refs(docx_path, output_dir, file_stem)
    print(f"  –ù–∞–π–¥–µ–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π: {len(image_map)}")

    # –®–∞–≥ 2: –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º –≤ HTML —á–µ—Ä–µ–∑ mammoth
    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –û–†–ò–ì–ò–ù–ê–õ–¨–ù–´–ô —Ñ–∞–π–ª, —á—Ç–æ–±—ã mammoth –æ—Å—Ç–∞–≤–∏–ª –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∫–∞–∫ rId (–Ω–µ –ø—Ä–µ–æ–±—Ä–∞–∑–æ–≤—ã–≤–∞–ª –≤ base64)
    print("üîÑ –ö–æ–Ω–≤–µ—Ä—Ç–∞—Ü–∏—è –≤ HTML...")
    
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value

    # –®–∞–≥ 3: –ó–∞–º–µ–Ω—è–µ–º —Å—Å—ã–ª–∫–∏ –Ω–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
    print("üîó –ü—Ä–æ–≤–µ—Ä–∫–∞ —Å—Å—ã–ª–æ–∫ –Ω–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è...")
    # –û—Ç–ª–∞–¥–æ—á–Ω—ã–π –≤—ã–≤–æ–¥: –ø—Ä–æ–≤–µ—Ä—è–µ–º, –∫–∞–∫–∏–µ img —Ç–µ–≥–∏ –µ—Å—Ç—å –≤ HTML
    soup_debug = BeautifulSoup(html, 'html.parser')
    img_tags = soup_debug.find_all('img')
    if img_tags:
        print(f"  üìã –ù–∞–π–¥–µ–Ω–æ img —Ç–µ–≥–æ–≤ –≤ HTML (–ø–µ—Ä–≤—ã–µ —Ç—Ä–∏ —Ç—ç–≥–∞): {len(img_tags)}")
        for i, img in enumerate(img_tags[:3]):  # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø–µ—Ä–≤—ã–µ 3 –¥–ª—è –æ—Ç–ª–∞–¥–∫–∏
            src = img.get('src', '')
            print(f"     img[{i}]: src='{src[:80]}...' (–ø–µ—Ä–≤—ã–µ 80 —Å–∏–º–≤–æ–ª–æ–≤)")
    
    # –ü–æ–ª—É—á–∞–µ–º –ø—É—Ç—å –∫ –ø–∞–ø–∫–µ —Å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è–º–∏
    images_dir = output_dir / images_folder_name

    
    html = replace_image_tags_in_html(html, image_map, images_folder_name, images_dir)
    
    # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç –∑–∞–º–µ–Ω—ã —Å—Å—ã–ª–æ–∫
    soup_after = BeautifulSoup(html, 'html.parser')
    img_tags_after = soup_after.find_all('img')
    if img_tags_after:
        print(f"  üìã –ü–æ—Å–ª–µ –∑–∞–º–µ–Ω—ã –Ω–∞–π–¥–µ–Ω–æ img —Ç–µ–≥–æ–≤: {len(img_tags_after)}")
        for i, img in enumerate(img_tags_after[:3]):  # –ü–æ–∫–∞–∑—ã–≤–∞–µ–º –ø–µ—Ä–≤—ã–µ 3 –¥–ª—è –æ—Ç–ª–∞–¥–∫–∏
            src = img.get('src', '')
            alt = img.get('alt', '')
            print(f"     img[{i}]: src='{src}', alt='{alt}'")

    # –®–∞–≥ 4: –ö–æ–Ω–≤–µ—Ä—Ç–∏—Ä—É–µ–º HTML ‚Üí Markdown
    print("üìù –ü—Ä–µ–æ–±—Ä–∞–∑–æ–≤–∞–Ω–∏–µ –≤ Markdown...")
    from markdownify import markdownify as md

    markdown_content = md(html, heading_style="ATX", strip=['style'])

    # –®–∞–≥ 4.5: –ò—Å–ø—Ä–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è (–∑–∞–º–µ–Ω—è–µ–º base64 –∏ <img> –Ω–∞ —Å—Å—ã–ª–∫–∏)
    print("üîß –ò—Å–ø—Ä–∞–≤–ª–µ–Ω–∏–µ —Å—Å—ã–ª–æ–∫ –Ω–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è...")
    markdown_content = fix_images_in_markdown(markdown_content, images_dir, images_folder_name)

    # –®–∞–≥ 4.6: –û—á–∏—â–∞–µ–º —Å–∫—Ä—ã—Ç—ã–µ –º–µ—Ç–∫–∏ –∏–∑ markdown
    print("üßπ –û—á–∏—Å—Ç–∫–∞ —Å–∫—Ä—ã—Ç—ã—Ö –º–µ—Ç–æ–∫...")
    markdown_content = clean_hidden_tags_in_markdown(markdown_content)



    # # –®–∞–≥ 5: –°–æ—Ö—Ä–∞–Ω—è–µ–º
    # with open(md_path, 'w', encoding='utf-8') as f:
    #     f.write(markdown_content)

    print(f"\n‚úÖ –ì–æ—Ç–æ–≤–æ!")
    # print(f"üìÑ Markdown: {md_path}")
    print(f"üñºÔ∏è  –ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è: {output_dir / images_folder_name}")
    
    return markdown_content

file_docx = r"X:\–£—á–µ–±–∞_–£–ò–ò\–ò—Ç–æ–≥–æ–≤—ã_–ü—Ä–æ–µ–∫—Ç\–≠—Ç–∞–ø ‚Ññ2.  AI_ML  –°–±–æ—Ä –±–∞–∑—ã\–ù–æ—Ä–º–∞—Ç–∏–≤–Ω–∞—è –±–∞–∑–∞\–ü–£–≠\DOCX\2.5.docx"
if __name__ == "__main__":
    sys_argv = None
    try:
        # –†–∞–±–æ—á–∞—è –≤–µ—Ä—Å–∏—è: –ø—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞–ª–∏—á–∏–µ –∞—Ä–≥—É–º–µ–Ω—Ç–∞ –∫–æ–º–∞–Ω–¥–Ω–æ–π —Å—Ç—Ä–æ–∫–∏
        if len(sys.argv) == 2 and sys.argv[1].strip():
            sys_argv = sys.argv[1]
        else:
            # –î–ª—è —Ç–µ—Å—Ç–æ–≤: –∏—Å–ø–æ–ª—å–∑—É–µ–º file_docx
            sys_argv = file_docx
        
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º —Å—É—â–µ—Å—Ç–≤–æ–≤–∞–Ω–∏–µ —Ñ–∞–π–ª–∞
        docx_path = Path(sys_argv)
        if not docx_path.exists():
            raise FileNotFoundError(f"–§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω: {sys_argv}")
        
        output_dir = docx_path.parent
        md_path = output_dir / f"{docx_path.stem}.md"

        markdown_content = docx_to_md_with_images(sys_argv)

        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        print(f"üìÑ Markdown: {md_path}")
    except FileNotFoundError as e:
        print(f"‚ùå –û—à–∏–±–∫–∞: {e}")
        print("–ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ: python docx_to_md_with_images.py <—Ñ–∞–π–ª.docx>")
        sys.exit(1)
    except Exception as e:
        print(f"‚ùå –û—à–∏–±–∫–∞: {e}")
        print("–ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ: python docx_to_md_with_images.py <—Ñ–∞–π–ª.docx>")
        sys.exit(1)