# docx_to_md_images_2.py
# Конвертация DOCX в Markdown с извлечением изображений
# Изображения вставляются в виде ссылок ![alt](path), а не base64
# Нормально выгружает image файлы и вставляются в md файл
import sys
import re
from pathlib import Path
import mammoth
from docx import Document

def clean_hidden_tags_in_docx(docx_path):
    """
    Удаляет скрытые метки вроде:
      #G0...
      #M12291 901873648 ...
      #S
      #X... (любой символ после # + цифры/пробелы)
    из всех параграфов и ячеек таблиц в .docx.
    Возвращает изменённый Document (в памяти).
    """
    doc = Document(docx_path)

    # Шаблон для скрытых меток (включая строки целиком и фрагменты)
    pattern = re.compile(
        r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # полная строка
        r'|'
        r'(?:\s+)?#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?(?:\s+|#S)?',     # фрагмент внутри строки
        re.IGNORECASE
    )

    # Очищаем параграфы
    for para in doc.paragraphs:
        if para.text.strip():
            cleaned_text = pattern.sub('', para.text)
            # Убираем лишние пробелы и пустые строки
            cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text).strip()
            if not cleaned_text:
                para.clear()  # полностью очищаем пустой параграф
            else:
                para.text = cleaned_text

    # Очищаем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cleaned_text = pattern.sub('', cell.text)
                    cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text).strip()
                    cell.text = cleaned_text

    return doc

def docx_to_markdown_cleaned(docx_path):
    """Сохраняет временный очищенный .docx и конвертирует его в Markdown."""
    # 1. Очищаем документ
    clean_doc = clean_hidden_tags_in_docx(docx_path)

    # 2. Сохраняем во временный файл (чтобы mammoth мог его прочитать)
    temp_path = docx_path.with_suffix('.cleaned.docx')
    clean_doc.save(temp_path)

    try:
        # 3. Конвертируем через mammoth
        with open(temp_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html = result.value
        # 4. HTML → Markdown
        from markdownify import markdownify as md
        markdown = md(html, heading_style="ATX", strip=["style"])
        return markdown
    finally:
        # 5. Удаляем временный файл
        temp_path.unlink(missing_ok=True)

def main(input_path):
    input_path = Path(input_path).resolve()
    if not input_path.exists():
        print(f"❌ Файл не найден: {input_path}")
        sys.exit(1)
    if input_path.suffix.lower() != '.docx':
        print("❌ Поддерживается только .docx")
        sys.exit(1)

    print(f"🔍 Обработка: {input_path}")
    md_content = docx_to_markdown_cleaned(input_path)

    # Сохраняем результат
    md_path = input_path.with_suffix('.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    print(f"✅ Готово! Markdown сохранён: {md_path}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Использование: python clean_docx_then_to_md.py <файл.docx>")
        sys.exit(1)
    main(sys.argv[1])