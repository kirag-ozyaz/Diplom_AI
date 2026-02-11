# Резервная копися первой верси. здесь картинки выкладываются в папки по имни обрабатываемого docx файла
import os
import sys
import base64
from pathlib import Path
import mammoth
from bs4 import BeautifulSoup
import re
import shutil
from zipfile import ZipFile
from docx import Document


def clean_hidden_tags_in_docx(docx_path):
    """
    Удаляет скрытые метки вроде:
      #G0...
      #M12291 901873648 ...
      #S
      #X... (любой символ после # + цифры/пробелы)
    из всех параграфов и ячеек таблиц в .docx.
    Возвращает изменённый Document (в памяти).
    """
    doc = Document(docx_path)

    # Шаблон для скрытых меток (включая строки целиком и фрагменты)
    # pattern = re.compile(
    #     r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # полная строка
    #     r'|'
    #     r'(?:\s+)?#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?(?:\s+|#S)?',     # фрагмент внутри строки
    #     re.IGNORECASE
    # )
    pattern = re.compile(
        r'^\s*(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z]|#S)\s*$'  # полная строка (оставляем как есть)
        r'|'
        r'(?:\s+)?(?:#[A-Z]\d+(?:\s+\d+(?:\s+\d+)*)?|#[A-Z])(?:\s+)?',  # ← ИЗМЕНЕНО: добавлен |#[A-Z]
        re.IGNORECASE
    )

    # Очищаем параграфы
    for para in doc.paragraphs:
        if para.text.strip():
            cleaned_text = pattern.sub('', para.text)
            # Убираем лишние пробелы и пустые строки
            cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text).strip()
            if not cleaned_text:
                para.clear()  # полностью очищаем пустой параграф
            else:
                para.text = cleaned_text

    # Очищаем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cleaned_text = pattern.sub('', cell.text)
                    cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text).strip()
                    cell.text = cleaned_text

    return doc


def extract_images_and_fix_refs(docx_path, output_dir, file_stem):
    """Извлекает изображения из .docx и возвращает словарь {rId: имя_файла}"""
    # Создаем папку с именем image_<имя_файла>
    images_dir = output_dir / f"image_{file_stem}"
    
    # Очищаем папку, если она существует
    if images_dir.exists():
        shutil.rmtree(images_dir)
    
    images_dir.mkdir(exist_ok=True)

    image_map = {}
    image_counter = 1

    # Извлекаем изображения через ZIP (т.к. .docx — это ZIP-архив)
    with ZipFile(docx_path, 'r') as docx_zip:
        # Сначала читаем relationships
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path not in docx_zip.namelist():
            print(f"  ⚠️  Файл {rels_path} не найден в архиве")
            return image_map
        
        rels_xml = docx_zip.read(rels_path).decode('utf-8')
        soup = BeautifulSoup(rels_xml, 'xml')
        
        # Получаем все Relationship элементы
        relationships = soup.find_all('Relationship')
        print(f"  📋 Найдено relationships: {len(relationships)}")
        
        # Парсим все связи изображений
        for rel in relationships:
            rel_type = rel.get('Type', '')
            # Проверяем, является ли это изображением
            if 'image' not in rel_type.lower():
                continue
            
            r_id = rel.get('Id', '')
            target = rel.get('Target', '')
            
            if not r_id or not target:
                print(f"  ⚠️  Пропущена связь: Id={r_id}, Target={target}")
                continue
            
            # Формируем полный путь к изображению
            # Target может быть: "media/image1.png" или "../media/image1.png" или просто "image1.png"
            if target.startswith('media/'):
                img_path_in_zip = f'word/{target}'
            elif target.startswith('../media/'):
                img_path_in_zip = f'word/{target[3:]}'  # Убираем ../
            elif '/' in target:
                # Если есть слэш, но не media/, пробуем как есть
                img_path_in_zip = f'word/{target}' if not target.startswith('word/') else target
            else:
                # Просто имя файла
                img_path_in_zip = f'word/media/{target}'
            
            # Проверяем, существует ли файл
            if img_path_in_zip not in docx_zip.namelist():
                # Пробуем альтернативные варианты
                alt_paths = [
                    f'word/media/{Path(target).name}',
                    f'word/{target}',
                    target
                ]
                found = False
                for alt_path in alt_paths:
                    if alt_path in docx_zip.namelist():
                        img_path_in_zip = alt_path
                        found = True
                        break
                
                if not found:
                    print(f"  ❌ Изображение не найдено: rId={r_id}, target={target}, пробовали: {img_path_in_zip}")
                    # Выводим список всех файлов в архиве для отладки
                    media_files = [f for f in docx_zip.namelist() if 'media' in f.lower() or 'image' in f.lower()]
                    if media_files:
                        print(f"     Доступные медиа-файлы: {media_files[:5]}...")  # Показываем первые 5
                    continue
            
            # Используем оригинальное имя файла из target
            # Извлекаем имя файла из пути (например, "media/image1.png" -> "image1.png")
            original_name = Path(target).name
            
            # Если имя файла пустое или нет расширения, используем счетчик
            if not original_name or not Path(original_name).suffix:
                ext = Path(target).suffix.lower()
                if not ext or ext not in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'):
                    ext = '.png'
                original_name = f"image{image_counter:03d}{ext}"
            
            # Проверяем, не существует ли уже файл с таким именем
            img_name = original_name
            img_path = images_dir / img_name
            counter = 1
            while img_path.exists():
                # Если файл уже существует, добавляем суффикс
                stem = Path(original_name).stem
                ext = Path(original_name).suffix
                img_name = f"{stem}_{counter}{ext}"
                img_path = images_dir / img_name
                counter += 1
            
            try:
                with open(img_path, 'wb') as f:
                    f.write(docx_zip.read(img_path_in_zip))
                
                image_map[r_id] = img_name
                print(f"  ✓ {r_id} → {img_name} (из {img_path_in_zip})")
                image_counter += 1
            except Exception as e:
                print(f"  ❌ Ошибка при сохранении изображения {r_id}: {e}")

    return image_map


def replace_image_tags_in_html(html, image_map, images_folder_name, images_dir, image_counter_start):
    """Заменяет <img> теги на корректные пути, сохраняя все изображения"""
    soup = BeautifulSoup(html, 'html.parser')
    image_counter = image_counter_start

    for img in soup.find_all('img'):
        src = img.get('src', '')
        new_src = None

        # 1. Обработка base64
        if src.startswith('data:image/'):
            try:
                match = re.match(r'data:image/(\w+);base64,(.+)', src)
                if match:
                    img_format = match.group(1).lower()
                    base64_data = match.group(2)
                    ext = {'png': '.png', 'jpeg': '.jpg', 'jpg': '.jpg', 'gif': '.gif', 'bmp': '.bmp',
                           'webp': '.webp'}.get(img_format, '.png')

                    img_data = base64.b64decode(base64_data)
                    import hashlib
                    img_hash = hashlib.md5(img_data).hexdigest()[:8]
                    img_name = f"img_{img_hash}{ext}"
                    img_path = images_dir / img_name

                    counter = 1
                    while img_path.exists():
                        stem = Path(img_name).stem
                        img_name = f"{stem}_{counter}{ext}"
                        img_path = images_dir / img_name
                        counter += 1

                    with open(img_path, 'wb') as f:
                        f.write(img_data)

                    new_src = f"{images_folder_name}/{img_name}"
                    print(f"  ✓ Извлечено из base64: {img_name}")
                    image_counter += 1
            except Exception as e:
                print(f"  ⚠️ Ошибка base64: {e}")

        # 2. Обработка rId
        elif src.startswith('rId') or (match := re.search(r'(rId\d+)', src)):
            r_id = match.group(1) if match else src
            if r_id in image_map:
                new_src = f"{images_folder_name}/{image_map[r_id]}"

        # 3. Если не распознано — НЕ удаляем, оставляем для отладки
        if new_src:
            img['src'] = new_src
            if not img.get('alt'):
                img['alt'] = Path(new_src).name
        else:
            # Оставляем тег, но помечаем для отладки
            print(f"  ⚠️ Необработанное изображение (src='{src[:60]}...'), оставлено как есть")

    return str(soup)


def fix_remaining_img_tags(markdown_text, images_folder_name):
    """Принудительно заменяет оставшиеся <img> теги в Markdown на синтаксис ![alt](src)"""
    soup = BeautifulSoup(markdown_text, 'html.parser')
    modified = False

    for img in soup.find_all('img'):
        src = img.get('src', '')
        alt = img.get('alt', Path(src).name if src else 'image')
        if src and not src.startswith(('http://', 'https://')):
            # Формируем корректный путь относительно MD-файла
            md_img = f'![{alt}]({src})'
            img.replace_with(md_img)
            modified = True

    if modified:
        # Конвертируем обратно в строку, удаляя оставшийся мусор от BeautifulSoup
        result = str(soup)
        # Удаляем обёртку <html><body> если есть
        result = re.sub(r'^<html><body>|</body></html>$', '', result, flags=re.IGNORECASE)
        return result
    return markdown_text

def docx_to_md_with_images(docx_path, output_dir=None):
    docx_path = Path(docx_path).resolve()
    if docx_path.suffix.lower() != '.docx':
        print("❌ Поддерживается только .docx")
        sys.exit(1)

    # Если output_dir не указан, используем папку с исходным файлом
    if output_dir is None:
        output_dir = docx_path.parent
    else:
        output_dir = Path(output_dir)
    
    # md_path = output_dir / f"{docx_path.stem}.md"

    # Шаг 0.1. Очищаем документ
    print("🖼  Очищаем документ...")
    clean_doc = clean_hidden_tags_in_docx(docx_path)
    # Шаг 0.2. Сохраняем во временный файл (чтобы mammoth мог его прочитать)
    temp_path = docx_path.with_suffix('.cleaned.docx')
    clean_doc.save(temp_path)


    # Шаг 1: Извлекаем изображения
    print("🖼️  Извлечение изображений...")
    file_stem = docx_path.stem
    images_folder_name = f"image_{file_stem}"
    image_map = extract_images_and_fix_refs(temp_path, output_dir, file_stem)
    print(f"  Найдено изображений: {len(image_map)}")

    # Шаг 2: Конвертируем в HTML через mammoth
    print("🔄 Конвертация в HTML...")
    
    with open(temp_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value

    # Шаг 3: Заменяем ссылки на изображения
    print("🔗 Проверка ссылок на изображения...")
    # Отладочный вывод: проверяем, какие img теги есть в HTML
    soup_debug = BeautifulSoup(html, 'html.parser')
    img_tags = soup_debug.find_all('img')
    if img_tags:
        print(f"  📋 Найдено img тегов в HTML: {len(img_tags)}")
        for i, img in enumerate(img_tags[:3]):  # Показываем первые 3 для отладки
            src = img.get('src', '')
            print(f"     img[{i}]: src='{src[:80]}...' (первые 80 символов)")
    
    # Получаем путь к папке с изображениями
    images_dir = output_dir / images_folder_name
    # Начальный счетчик для новых изображений (продолжаем после уже извлеченных)
    image_counter_start = len(image_map) + 1
    
    html = replace_image_tags_in_html(html, image_map, images_folder_name, images_dir, image_counter_start)
    
    # Проверяем результат замены ссылок
    soup_after = BeautifulSoup(html, 'html.parser')
    img_tags_after = soup_after.find_all('img')
    if img_tags_after:
        print(f"  📋 После замены найдено img тегов: {len(img_tags_after)}")
        for i, img in enumerate(img_tags_after[:3]):  # Показываем первые 3 для отладки
            src = img.get('src', '')
            alt = img.get('alt', '')
            print(f"     img[{i}]: src='{src}', alt='{alt}'")

    # Шаг 4: Конвертируем HTML → Markdown
    print("📝 Преобразование в Markdown...")
    from markdownify import markdownify as md

    markdown_content = md(html, heading_style="ATX", strip=['style'])

    # Принудительная пост-обработка изображений
    markdown_content = fix_remaining_img_tags(markdown_content, images_folder_name)

    # Дополнительная очистка: замена оставшихся <img> тегов через регулярку
    def replace_img_tags(text):
        def replacer(m):
            src = re.search(r'src="([^"]+)"', m.group(0))
            alt = re.search(r'alt="([^"]+)"', m.group(0))
            src_val = src.group(1) if src else ''
            alt_val = alt.group(1) if alt else Path(src_val).name if src_val else 'image'
            if src_val and not src_val.startswith(('http://', 'https://')):
                return f'\n\n![{alt_val}]({src_val})\n\n'
            return m.group(0)

        return re.sub(r'<img[^>]+>', replacer, text, flags=re.IGNORECASE)

    markdown_content = replace_img_tags(markdown_content)

    # Проверка результата
    img_pattern = r'!\[.*?\]\(.*?\)'
    img_matches = re.findall(img_pattern, markdown_content)
    if img_matches:
        print(f"  ✓ Найдено ссылок на изображения в Markdown: {len(img_matches)}")
        for i, match in enumerate(img_matches[:5]):  # Показываем первые 5
            print(f"     {i + 1}: {match}")
    else:
        print(f"  ⚠️  Ссылки на изображения не найдены в Markdown!")
        # Отладка: показываем фрагмент markdown
        print(f"  📋 Первые 500 символов:\n{markdown_content[:500]}")



    # # Шаг 5: Сохраняем
    # with open(md_path, 'w', encoding='utf-8') as f:
    #     f.write(markdown_content)

    # Шаг 5. Удаляем временный файл
    temp_path.unlink(missing_ok=True)

    print(f"\n✅ Готово!")
    # print(f"📄 Markdown: {md_path}")
    print(f"🖼️  Изображения: {output_dir / images_folder_name}")
    
    return markdown_content

file_docx = r"X:\Учеба_УИИ\Итоговы_Проект\Этап №2.  AI_ML  Сбор базы\Нормативная база\ПУЭ\DOCX\2.5.docx"
if __name__ == "__main__":
    sys_argv = None
    try:
        # Рабочая версия: проверяем наличие аргумента командной строки
        if len(sys.argv) == 2 and sys.argv[1].strip():
            sys_argv = sys.argv[1]
        else:
            # Для тестов: используем file_docx
            sys_argv = file_docx
        
        # Проверяем существование файла
        docx_path = Path(sys_argv)
        if not docx_path.exists():
            raise FileNotFoundError(f"Файл не найден: {sys_argv}")
        
        output_dir = docx_path.parent
        md_path = output_dir / f"{docx_path.stem}.md"

        markdown_content = docx_to_md_with_images(sys_argv)

        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        print(f"📄 Markdown: {md_path}")
    except FileNotFoundError as e:
        print(f"❌ Ошибка: {e}")
        print("Использование: python docx_to_md_with_images.py <файл.docx>")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        print("Использование: python docx_to_md_with_images.py <файл.docx>")
        sys.exit(1)